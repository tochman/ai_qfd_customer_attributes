from models import PrimaryAttribute, BusinessAnalysis
from docx import Document
from docx.shared import Inches
import logging
import os
import markdown
from bs4 import BeautifulSoup
from typing import List, Dict

logger = logging.getLogger()


def add_markdown_to_docx(markdown_text: str, doc: Document):
    """
    Convert markdown text to HTML and add it to the DOCX document.
    """
    # Convert Markdown to HTML
    html = markdown.markdown(markdown_text)

    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    # Iterate through parsed content and add it to the document
    for element in soup:
        if element.name == "h1":
            doc.add_heading(element.get_text(), level=1)
        elif element.name == "h2":
            doc.add_heading(element.get_text(), level=2)
        elif element.name == "h3":
            doc.add_heading(element.get_text(), level=3)
        elif element.name == "p":
            doc.add_paragraph(element.get_text())
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif element.name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Number")
        elif element.name == "strong":
            # Add bold text in a paragraph
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.bold = True
        elif element.name == "em":
            # Add italic text in a paragraph
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.italic = True
        # Add other tags as needed...


def save_attributes_json(
    customer_attributes: List[PrimaryAttribute], file_path: str = "./attributes.json"
) -> None:
    """
    Save the derived customer attributes to a JSON file for manual review.
    """
    from models import CustomerAttributes

    logger.info(f"Saving customer attributes to {file_path} for manual review...")
    try:
        # Serialize the Pydantic model to JSON using model_dump_json()
        attributes_json = CustomerAttributes(
            attributes=customer_attributes
        ).model_dump_json(indent=4)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(attributes_json)
        logger.info(f"Customer attributes successfully saved to {file_path}")
    except Exception as e:
        logger.error(f"Failed to save customer attributes to {file_path}: {e}")


def generate_final_report(
    domain: str,
    customer_attributes: List[PrimaryAttribute],
    attribute_weights: Dict[str, float],
    business_analysis: BusinessAnalysis,
) -> None:
    """
    Generate a comprehensive DOCX report of the analysis.
    """
    logger.info("Generating final report...")

    # Create DOCX document
    doc = Document()
    # Use the title from the BusinessAnalysis object
    doc.add_heading(business_analysis.title, 0)

    # Add Introduction
    doc.add_heading("Introduction", level=1)
    add_markdown_to_docx(business_analysis.introduction, doc)

    # Add Derived Attributes Introduction
    doc.add_heading("Introduction to Derived Customer Attributes", level=1)
    add_markdown_to_docx(business_analysis.derived_attributes_introduction, doc)

    # Add Findings
    doc.add_heading("Analysis", level=1)
    add_markdown_to_docx(business_analysis.analysis, doc)

    # Add Recommendations Section
    doc.add_heading("Recommendations", level=1)
    add_markdown_to_docx(business_analysis.recommendations, doc)

    # Add Conclusion
    doc.add_heading("Conclusion", level=1)
    add_markdown_to_docx(business_analysis.conclusions, doc)

    # Add Customer Attribute Breakdown Section
    doc.add_heading("Customer Attribute Breakdown", level=1)

    # Create the customer attribute table
    attribute_table = doc.add_table(rows=1, cols=3)
    attribute_table.style = "Light List Accent 1"

    # Set the header cells
    hdr_cells = attribute_table.rows[0].cells
    hdr_cells[0].text = "Primary Attribute"
    hdr_cells[1].text = "Secondary Attribute"
    hdr_cells[2].text = "Tertiary Attribute"

    # Iterate over the customer attributes and fill the table
    for attr in customer_attributes:
        primary_attr_written = (
            False  # Flag to track if primary attribute has been written
        )
        for secondary_attr in attr.secondary_attributes:
            secondary_attr_written = False  # Flag for secondary attribute
            for tertiary_attr in secondary_attr.tertiary_attributes:
                # Add a new row to the table
                row_cells = attribute_table.add_row().cells

                # Only write the primary attribute if it hasn't been written yet
                if not primary_attr_written:
                    row_cells[0].text = attr.primary_attribute
                    primary_attr_written = True
                else:
                    row_cells[0].text = ""  # Leave empty for subsequent rows

                # Only write the secondary attribute if it hasn't been written yet
                if not secondary_attr_written:
                    row_cells[1].text = secondary_attr.attribute
                    secondary_attr_written = True
                else:
                    row_cells[1].text = ""  # Leave empty for subsequent rows

                # Write the tertiary attribute
                row_cells[2].text = tertiary_attr.attribute

            # Reset the secondary attribute flag after finishing tertiary attributes
            secondary_attr_written = False

        # Reset the primary attribute flag after finishing secondary attributes
        primary_attr_written = False

    # Add spacing after the table
    doc.add_paragraph()

    # Relative Importance Table
    doc.add_heading("Relative Importance of Customer Attributes", level=1)
    importance_table = doc.add_table(rows=1, cols=2)
    importance_table.style = "Light List Accent 1"
    hdr_cells = importance_table.rows[0].cells
    hdr_cells[0].text = "Customer Attribute"
    hdr_cells[1].text = "Relative Importance (%)"

    for attr, importance in attribute_weights.items():
        row_cells = importance_table.add_row().cells
        row_cells[0].text = attr
        row_cells[1].text = f"{importance:.2f}%"

    # Add spacing after the table
    doc.add_paragraph()

    # Add Chart
    chart_filename = "relative_importance_chart.png"
    if os.path.exists(chart_filename):
        doc.add_heading("Relative Importance Chart", level=1)
        doc.add_picture(chart_filename, width=Inches(6))
    else:
        logger.warning(
            f"Chart file {chart_filename} not found. Skipping chart insertion."
        )

    # Add Appendix Section: Statements grouped by sentiment
    doc.add_heading("Appendix: Statements by Sentiment", level=1)

    # Iterate over primary attributes and create the appendix table
    appendix_table = doc.add_table(rows=1, cols=4)
    appendix_table.style = "Light List Accent 1"

    # Set the header cells
    appendix_hdr_cells = appendix_table.rows[0].cells
    appendix_hdr_cells[0].text = "Primary Attribute"
    appendix_hdr_cells[1].text = "Positive Statements"
    appendix_hdr_cells[2].text = "Neutral Statements"
    appendix_hdr_cells[3].text = "Negative Statements"

    # Group statements by primary attribute and sentiment
    for primary_attr in customer_attributes:
        positive_statements = []
        neutral_statements = []
        negative_statements = []

        for secondary_attr in primary_attr.secondary_attributes:
            for tertiary_attr in secondary_attr.tertiary_attributes:
                for stmt in tertiary_attr.statements:
                    sentiment = stmt.score
                    if sentiment >= 0.3:
                        positive_statements.append(
                            f"{stmt.statement} ({sentiment:.2f})"
                        )
                    elif sentiment <= -0.3:
                        negative_statements.append(
                            f"{stmt.statement} ({sentiment:.2f})"
                        )
                    else:
                        neutral_statements.append(f"{stmt.statement} ({sentiment:.2f})")

        # Add a new row for the primary attribute
        row_cells = appendix_table.add_row().cells
        row_cells[0].text = primary_attr.primary_attribute
        row_cells[1].text = (
            "\n".join(positive_statements) if positive_statements else "-"
        )
        row_cells[2].text = "\n".join(neutral_statements) if neutral_statements else "-"
        row_cells[3].text = (
            "\n".join(negative_statements) if negative_statements else "-"
        )

    # Save DOCX file
    report_filename = "final_customer_feedback_report.docx"
    try:
        doc.save(report_filename)
        logger.info(f"Final report saved to {report_filename}")
    except Exception as e:
        logger.error(f"An error occurred while saving the report: {e}")
