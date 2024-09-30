import os
import json
import logging
import argparse
import re  # Importing the re module
import pdb
from typing import List, Dict
from dotenv import load_dotenv
from pydantic import BaseModel, Field
import matplotlib.pyplot as plt
import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt

# -----------------------------------------------------
# LangChain Imports
# -----------------------------------------------------
from langchain.prompts import ChatPromptTemplate
from langchain.output_parsers import PydanticOutputParser
from langchain_core.runnables import RunnableSequence
from langchain_openai import ChatOpenAI
from langchain_core.exceptions import OutputParserException

# -----------------------------------------------------
# NLTK Imports for Preprocessing
# -----------------------------------------------------
import nltk
from nltk.stem import WordNetLemmatizer

# -----------------------------------------------------
# Download Necessary NLTK Data
# -----------------------------------------------------
nltk.download("wordnet")
nltk.download("omw-1.4")  # For lemmatization

lemmatizer = WordNetLemmatizer()


# -----------------------------------------------------
# Logging Configuration
# -----------------------------------------------------
def configure_logging():
    """
    Configure logging with both console and file handlers.
    """
    logger = logging.getLogger()
    logger.setLevel(logging.DEBUG)  # Capture all levels for file

    # Create formatter
    formatter = logging.Formatter("%(asctime)s - %(levelname)s - %(message)s")

    # Console handler for INFO level and above
    ch = logging.StreamHandler()
    ch.setLevel(logging.INFO)
    ch.setFormatter(formatter)
    logger.addHandler(ch)

    # File handler for DEBUG level
    fh = logging.FileHandler("debug.log")
    fh.setLevel(logging.DEBUG)
    fh.setFormatter(formatter)
    logger.addHandler(fh)

    return logger


logger = configure_logging()

# -----------------------------------------------------
# Environment Variables and Pydantic Models
# -----------------------------------------------------
# Load environment variables from .env file
load_dotenv()

# Access the API key
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    logger.error("OpenAI API key not found. Please set it in your .env file.")
    raise ValueError("OpenAI API key not found. Please set it in your .env file.")


# Define the Pydantic schema for the expected output
class Statement(BaseModel):
    """Model representing a customer statement with its sentiment score."""

    statement: str = Field(description="The original customer statement")
    sentiment_score: float = Field(description="Sentiment score between -1 and 1")


class TertiaryAttribute(BaseModel):
    """Model representing a tertiary attribute."""

    attribute: str = Field(description="The name of the tertiary attribute")
    statements: List[Statement] = Field(
        description="List of customer statements related to this attribute along with their sentiment scores"
    )


class SecondaryAttribute(BaseModel):
    """Model representing a secondary attribute."""

    attribute: str = Field(description="The name of the secondary attribute")
    tertiary_attributes: List[TertiaryAttribute] = Field(
        description="List of tertiary attributes under this secondary attribute"
    )


class PrimaryAttribute(BaseModel):
    """Model representing a primary attribute."""

    primary_attribute: str = Field(description="The name of the primary attribute")
    secondary_attributes: List[SecondaryAttribute] = Field(
        description="List of secondary attributes under this primary attribute"
    )


class CustomerAttributes(BaseModel):
    """Model representing a list of customer attributes."""

    attributes: List[PrimaryAttribute] = Field(
        description="List of primary customer attributes with nested secondary and tertiary attributes"
    )


class SentimentResult(BaseModel):
    """Model representing the sentiment analysis result for a statement."""

    statement: str = Field(description="The original customer statement")
    score: float = Field(description="Sentiment score between -1 and 1")
    label: str = Field(description="Sentiment label: Positive, Neutral, or Negative")


class SentimentResults(BaseModel):
    """Model representing a list of sentiment analysis results."""

    results: List[SentimentResult] = Field(description="List of sentiment results")


class BusinessAnalysis(BaseModel):
    """Model representing the business analysis based on previous results"""

    introduction: str = Field(
        description="The introduction of the narrative to the analysis"
    )
    analysis: str = Field(description="The main body of the narrative of the analysis")
    recommendations: str = Field(
        description="Summary and recommendations for the operations managment team"
    )
    derived_attributes_introduction: str = Field(
        description="The introduction of the process of deriving customer attributes from survey statements"
    )
    conclusions: str = Field(description="Summary and conclusion of the analysis")


llm = ChatOpenAI(model_name="gpt-4o-mini", temperature=0.3, openai_api_key=openai_api_key)


# -----------------------------------------------------
# Argument Parsing
# -----------------------------------------------------
def parse_arguments():
    """
    Parse command-line arguments.
    """
    parser = argparse.ArgumentParser(
        description="Customer Feedback Analysis with Nested Attributes"
    )
    parser.add_argument(
        "--domain", type=str, default="Healthcare Services", help="Domain for analysis"
    )
    parser.add_argument(
        "--input_file",
        type=str,
        default="./survey.txt",
        help="Path to customer statements file",
    )
    parser.add_argument(
        "--log_level",
        type=str,
        default="INFO",
        choices=["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"],
        help="Set the logging level",
    )
    return parser.parse_args()


# -----------------------------------------------------
# Loading Customer Statements
# -----------------------------------------------------
def load_customer_statements(file_path: str = "./survey.txt") -> List[str]:
    """
    Load customer statements from a text file.

    Args:
        file_path (str): Path to the text file containing customer statements.

    Returns:
        List[str]: A list of customer statements.
    """
    try:
        logger.info(f"Loading customer statements from {file_path}...")
        with open(file_path, "r", encoding="utf-8") as file:
            statements = [line.strip() for line in file if line.strip()]
        logger.debug(f"Loaded {len(statements)} customer statements.")
        return statements
    except FileNotFoundError:
        logger.error(f"File not found: {file_path}")
        return []
    except Exception as e:
        logger.error(f"An error occurred while loading customer statements: {e}")
        return []


# -----------------------------------------------------
# Loading Survey Details
# -----------------------------------------------------
def load_survey_details(file_path: str = "./survey_details.txt") -> str:
    """
    Load survey details from a text file.

    Args:
        file_path (str): Path to the text file containing survey details.

    Returns:
        str: A string containing survey details for additional context.
    """
    try:
        logger.info(f"Loading survey details from {file_path}...")
        with open(file_path, "r", encoding="utf-8") as file:
            survey_details = file.read().strip()
        logger.debug(f"Loaded survey details: {survey_details}")
        return survey_details
    except FileNotFoundError:
        logger.error(f"File not found: {file_path}")
        return ""
    except Exception as e:
        logger.error(f"An error occurred while loading survey details: {e}")
        return ""



# -----------------------------------------------------
# Preprocessing Function
# -----------------------------------------------------
def preprocess_text(text: str) -> str:
    """
    Preprocess text by lowercasing, removing punctuation, and lemmatizing.

    Args:
        text (str): The text to preprocess.

    Returns:
        str: The preprocessed text.
    """
    text = text.lower()
    text = re.sub(r"[^\w\s]", "", text)
    words = text.split()
    words = [lemmatizer.lemmatize(word) for word in words]
    return " ".join(words)


# -----------------------------------------------------
# Creating the Sentiment Analysis Chain
# -----------------------------------------------------
def create_sentiment_chain() -> RunnableSequence:
    """
    Create a sentiment analysis chain using LangChain's ChatPromptTemplate and PydanticOutputParser.

    Returns:
        RunnableSequence: A sequence that performs sentiment analysis and returns structured results.
    """
    logger.info("Creating sentiment analysis chain...")

    # Initialize the Pydantic output parser and retrieve format instructions
    sentiment_parser = PydanticOutputParser(pydantic_object=SentimentResults)
    format_instructions = sentiment_parser.get_format_instructions()

    # Use format instructions in the prompt, treating it as a partial variable
    sentiment_prompt_template = ChatPromptTemplate.from_template(
        """
    Perform sentiment analysis on the following customer statements in the {domain} domain.
    For each statement, provide a sentiment score between -1 (negative) and +1 (positive),
    and a sentiment label (Positive, Neutral, Negative).

    **Important:** Ensure that each result includes the 'statement', 'score', and 'label' fields,
    even if the input statement is incomplete or unclear.

    Please return the results in the following format:
    {format_instructions}

    Customer Statements:
    {statements}
        """
    ).partial(format_instructions=format_instructions)

    # Define the sequence
    sentiment_sequence = RunnableSequence(
        sentiment_prompt_template,
        llm,
        sentiment_parser,  # Parse the output into SentimentResults
    )

    return sentiment_sequence


# -----------------------------------------------------
# Creating the Attribute Derivation Chain
# -----------------------------------------------------
def create_attribute_chain() -> RunnableSequence:
    """
    Create an attribute derivation chain with a nested attribute structure.

    Returns:
        RunnableSequence: A sequence that derives customer attributes from sentiment analysis.
    """
    logger.info("Creating attribute derivation chain...")

    # Initialize the Pydantic output parser and retrieve format instructions
    attribute_parser = PydanticOutputParser(pydantic_object=CustomerAttributes)
    format_instructions = attribute_parser.get_format_instructions()

    # Use format instructions in the prompt, treating it as a partial variable
    attribute_prompt_template = ChatPromptTemplate.from_template(
        """
    Based on the statements and sentiment analysis provided, identify key customer attributes and critical features for 
    the {domain} domain using the Quality function deployment approach to ensure that customer needs and expectations are 
    incorporated into the final design of the service. 
    Structure the attributes in a three-level hierarchy (primary, secondary, tertiary) with associated customer statements and their sentiment scores.

    1. **Primary Attributes**: Broad categories of customer needs.
    2. **Secondary Attributes**: Specific aspects within each primary attribute.
    3. **Tertiary Attributes**: Detailed elements that impact each secondary attribute, along with related customer statements and their sentiment scores.

    Ensure that:
    - Attribute names are unique within their respective levels.
    - Each primary attribute can have multiple secondary attributes.
    - Each secondary attributes can have multiple tertiary attributes
    - Each customer statement is assigned to one and only one tertiary attribute.
    - The structure is comprehensive and professionally worded relevant to the {domain} domain.
    - no statements are left out without being attributed to an attribute.

    **Example Output:**
    
    ```json
    {{
        "attributes": [
            {{
                "primary_attribute": "Staff Interaction",
                "secondary_attributes": [
                    {{
                        "attribute": "Staff Professionalism",
                        "tertiary_attributes": [
                            {{
                                "attribute": "Politeness of staff",
                                "statements": [
                                    {{
                                        "statement": "The staff was extremely polite and accommodating.",
                                        "sentiment_score": 0.95
                                    }},
                                    {{
                                        "statement": "Great service and professional staff.",
                                        "sentiment_score": 0.9
                                    }}
                                ]
                            }},
                            {{
                                "attribute": "Communication skills of staff",
                                "statements": [
                                    {{
                                        "statement": "I appreciated the clear communication from the staff.",
                                        "sentiment_score": 0.85
                                    }},
                                    {{
                                        "statement": "I don't understand the nurse. Language training please!",
                                        "sentiment_score": -0.7
                                    }}
                                ]
                            }}
                        ]
                    }}
                ]
            }}
            // Add more primary attributes as needed
        ]
    }}
    ```
    
    Please return the results in the following format:
    {format_instructions}
    
    Sentiment Analysis Results:
    {sentiment_results}
        """
    ).partial(format_instructions=format_instructions)
    # Define the sequence
    attribute_sequence = RunnableSequence(
        attribute_prompt_template,
        llm,
        attribute_parser,  # Parse the output into CustomerAttributes
    )

    return attribute_sequence


# -----------------------------------------------------
# Calculating Relative Importance
# -----------------------------------------------------
def calculate_relative_importance(
    customer_attributes: List[PrimaryAttribute]
) -> Dict[str, float]:
    """
    Calculate the relative importance of customer attributes based on sentiment scores.

    Args:
        customer_attributes (List[PrimaryAttribute]): List of customer attributes derived from sentiment analysis.

    Returns:
        Dict[str, float]: A dictionary mapping primary attributes to their relative importance percentages.
    """
    logger.info("Calculating relative importance of customer attributes...")
    attribute_weights = {}
    total_score = 0.0

    # Initialize weights for primary attributes
    for attr in customer_attributes:
        primary_attr = attr.primary_attribute
        attribute_weights[primary_attr] = 0.0

    # Aggregate absolute sentiment scores to avoid negative weights
    for primary_attr in customer_attributes:
        for secondary_attr in primary_attr.secondary_attributes:
            for tertiary_attr in secondary_attr.tertiary_attributes:
                for stmt in tertiary_attr.statements:
                    score = abs(
                        stmt.sentiment_score
                    )  # Take absolute value of the score
                    attribute_weights[primary_attr.primary_attribute] += score
                    total_score += score
                    logger.debug(
                        f"Adding absolute score {score:.2f} from statement '{stmt.statement}' to attribute '{primary_attr.primary_attribute}'"
                    )

    if total_score == 0:
        logger.error(
            "Total sentiment score is zero, cannot calculate attribute weights."
        )
        return attribute_weights

    # Convert to percentages and normalize so that the sum is 100%
    for attr in attribute_weights:
        attribute_weights[attr] = round(
            (attribute_weights[attr] / total_score) * 100, 2
        )

    # Calculate total to verify correctness
    total_weight = sum(attribute_weights.values())
    logger.info(f"Total weight calculated (should be ~100%): {total_weight}%")

    return attribute_weights


# -----------------------------------------------------
# Generating Visualizations
# -----------------------------------------------------
def generate_visualizations(attribute_weights: Dict[str, float]) -> None:
    """
    Generate and display a bar chart of attribute weights.

    Args:
        attribute_weights (Dict[str, float]): Dictionary of attribute weights.
    """
    logger.info("Generating visualizations...")

    if not attribute_weights:
        logger.error("Attribute weights are empty, skipping visualization.")
        return

    primary_attrs = list(attribute_weights.keys())
    weights = list(attribute_weights.values())

    # Enhanced Visualization using Seaborn for better aesthetics
    import seaborn as sns  # Imported here to avoid issues if Seaborn is not used elsewhere

    sns.set(style="whitegrid")
    plt.figure(figsize=(max(12, len(primary_attrs) * 1.5), 8))

    # Remove hue parameter to fix FutureWarning
    barplot = sns.barplot(x=primary_attrs, y=weights, palette="Blues_d")
    plt.xlabel("Primary Attributes", fontsize=12)
    plt.ylabel("Total Weight (%)", fontsize=12)
    plt.title("Relative Importance of Primary Customer Attributes", fontsize=16)
    plt.xticks(rotation=45, ha="right", fontsize=10)
    plt.yticks(fontsize=10)

    # Annotate bars with weights
    for index, bar in enumerate(barplot.patches):
        height = bar.get_height()
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            height + 0.5,
            f"{weights[index]:.2f}%",
            ha="center",
            va="bottom",
            fontsize=10,
        )

    plt.tight_layout()

    # Save chart
    chart_filename = "relative_importance_chart.png"
    plt.savefig(chart_filename, dpi=300)
    logger.info(f"Chart saved to {chart_filename}")
    # plt.show()


# -----------------------------------------------------
# Generating Business Analysis
# -----------------------------------------------------
def generate_business_analysis(
    domain: str,
    customer_attributes: List[PrimaryAttribute],
    attribute_weights: Dict[str, float],
    survey_details: str,
) -> BusinessAnalysis:
    """
    Generate a comprehensive business analysis based on the sentiment analysis, customer attributes, and survey details.

    Args:
        domain (str): The domain of the analysis.
        customer_attributes (List[PrimaryAttribute]): List of customer attributes.
        attribute_weights (Dict[str, float]): Dictionary of attribute weights.
        survey_details (str): Unstructured text containing survey details for additional context.

    Returns:
        BusinessAnalysis: A Pydantic object containing the business analysis.
    """
    logger.info("Generating business analysis using LLM...")

    # Prepare the data in a structured format
    attribute_data = []
    for primary_attr in customer_attributes:
        for secondary_attr in primary_attr.secondary_attributes:
            for tertiary_attr in secondary_attr.tertiary_attributes:
                for stmt in tertiary_attr.statements:
                    attribute_data.append(
                        {
                            "Primary Attribute": primary_attr.primary_attribute,
                            "Secondary Attribute": secondary_attr.attribute,
                            "Tertiary Attribute": tertiary_attr.attribute,
                            "Statement": stmt.statement,
                            "Score": stmt.sentiment_score,
                        }
                    )

    # Serialize the data to JSON for inclusion in the prompt
    attribute_json = json.dumps(attribute_data, indent=2, ensure_ascii=False)
    weights_json = json.dumps(attribute_weights, indent=2, ensure_ascii=False)

    # Create a prompt for the LLM to generate the analysis
    prompt = f"""
You are a professional business analyst specialized in Quality function deployment. Below are the details of a survey that was conducted, followed by the customer feedback analysis for the {domain} domain. Provide a comprehensive business analysis suitable for inclusion in a report for operations management. Use the context provided in the survey details to customize your analysis.

### Survey Details
{survey_details}

### Customer Attribute Breakdown
{attribute_json}

### Relative Importance of Customer Attributes
{weights_json}

Please structure your analysis into five sections:
1. **Introduction**: Provide a background of the report, what it is based on, who it is for, and who commissioned it/led the project together with an overview of the business analysis using the details about the survey and the generally accepted standards for the {domain}.
2. **Introduction to Derived Customer Attributes**: Based on techniques used in QFD, provide an explanation on how customer attributes can be derived from a set of statements from the customers.
3. **Analysis**: Discuss key findings and insights derived from the data.
4. **Recommendations**: Offer actionable suggestions for the operations management based on the analysis.
5. **Conclusion**: A high-level conclusion of the analysis with a focus on customer satisfaction and quality.

Ensure the analysis is professional, clear, and suitable for a business report aimed at the operational management team.
"""

    # Initialize the Pydantic output parser for BusinessAnalysis
    business_analysis_parser = PydanticOutputParser(pydantic_object=BusinessAnalysis)
    business_analysis_format_instructions = (
        business_analysis_parser.get_format_instructions()
    )

    # Append format instructions to the prompt
    full_prompt = prompt + "\n\n" + business_analysis_format_instructions

    # Generate the analysis using invoke() instead of __call__()
    try:
        response = llm.invoke(full_prompt)
        # Parse the response into BusinessAnalysis
        business_analysis = business_analysis_parser.parse(response.content)
        logger.debug(f"Business Analysis Generated: {business_analysis}")
        return business_analysis
    except OutputParserException as e:
        logger.error(f"Failed to parse business analysis: {e}")
        logger.debug(f"Assistant's response content: {response.content}")
        return BusinessAnalysis(
            introduction="Business analysis could not be generated due to a parsing error.",
            derived_attributes_introduction="",
            analysis="",
            recommendations="",
            conclusions="",
        )
    except Exception as e:
        logger.error(f"Failed to generate business analysis: {e}")
        return BusinessAnalysis(
            introduction="Business analysis could not be generated due to an error.",
            derived_attributes_introduction="",
            analysis="",
            recommendations="",
            conclusions="",
        )


# -----------------------------------------------------
# Utility Function to Add Markdown Content to DOCX
# -----------------------------------------------------
def add_markdown_to_docx(markdown_text: str, doc: Document):
    """
    Convert markdown text to HTML and add it to the DOCX document.

    Args:
        markdown_text (str): The markdown text to add.
        doc (Document): The docx document object.
    """
    # Convert Markdown to HTML
    html = markdown.markdown(markdown_text)

    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html, "html.parser")

    # Iterate through parsed content and add it to the document
    for element in soup:
        if element.name == "h1":
            doc.add_heading(element.get_text(), level=1)
        elif element.name == "h2":
            doc.add_heading(element.get_text(), level=2)
        elif element.name == "h3":
            doc.add_heading(element.get_text(), level=3)
        elif element.name == "p":
            doc.add_paragraph(element.get_text())
        elif element.name == "ul":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Bullet")
        elif element.name == "ol":
            for li in element.find_all("li"):
                doc.add_paragraph(li.get_text(), style="List Number")
        elif element.name == "strong":
            # Add bold text in a paragraph
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.bold = True
        elif element.name == "em":
            # Add italic text in a paragraph
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.italic = True
        # Add other tags as needed...


# -----------------------------------------------------
# Generating the Final Report
# -----------------------------------------------------
def generate_final_report(
    domain: str,
    customer_attributes: List[PrimaryAttribute],
    attribute_weights: Dict[str, float],
    business_analysis: BusinessAnalysis,
) -> None:
    """
    Generate a comprehensive DOCX report of the analysis.

    Args:
        domain (str): The domain of the analysis.
        customer_attributes (List[PrimaryAttribute]): List of customer attributes.
        attribute_weights (Dict[str, float]): Dictionary of attribute weights.
        business_analysis (BusinessAnalysis): The business analysis object.
    """
    logger.info("Generating final report...")

    # Create DOCX document
    doc = Document()
    doc.add_heading(f"Customer Feedback Report for {domain}", 0)

    # Add Introduction
    doc.add_heading("Introduction", level=1)
    add_markdown_to_docx(business_analysis.introduction, doc)

    # Add Findings
    doc.add_heading("Findings", level=1)

    # Customer Attribute Breakdown Section
    doc.add_heading("Customer Attribute Breakdown", level=2)
    add_markdown_to_docx(business_analysis.derived_attributes_introduction, doc)

    # Create the customer attribute table
    attribute_table = doc.add_table(rows=1, cols=3)
    attribute_table.style = "Light List Accent 1"

    # Set the header cells
    hdr_cells = attribute_table.rows[0].cells
    hdr_cells[0].text = "Primary Attribute"
    hdr_cells[1].text = "Secondary Attribute"
    hdr_cells[2].text = "Tertiary Attribute"

    # Iterate over the customer attributes and fill the table
    for attr in customer_attributes:
        primary_attr_written = (
            False  # Flag to track if primary attribute has been written
        )
        for secondary_attr in attr.secondary_attributes:
            secondary_attr_written = (
                False  # Flag to track if secondary attribute has been written
            )
            for tertiary_attr in secondary_attr.tertiary_attributes:
                # Add a new row to the table
                row_cells = attribute_table.add_row().cells

                # Only write the primary attribute if it hasn't been written yet
                if not primary_attr_written:
                    row_cells[0].text = attr.primary_attribute
                    primary_attr_written = True
                else:
                    row_cells[0].text = ""  # Leave empty for subsequent rows

                # Only write the secondary attribute if it hasn't been written yet
                if not secondary_attr_written:
                    row_cells[1].text = secondary_attr.attribute
                    secondary_attr_written = True
                else:
                    row_cells[1].text = ""  # Leave empty for subsequent rows

                # Write the tertiary attribute
                row_cells[2].text = tertiary_attr.attribute

    # Add spacing after the table
    doc.add_paragraph()

    # Relative Importance Table
    doc.add_heading("Relative Importance of Customer Attributes", level=2)
    importance_table = doc.add_table(rows=1, cols=2)
    importance_table.style = "Light List Accent 1"
    hdr_cells = importance_table.rows[0].cells
    hdr_cells[0].text = "Customer Attribute"
    hdr_cells[1].text = "Relative Importance (%)"

    for attr, importance in attribute_weights.items():
        row_cells = importance_table.add_row().cells
        row_cells[0].text = attr
        row_cells[1].text = f"{importance:.2f}%"

    # Add spacing after the table
    doc.add_paragraph()

    # Add Chart
    chart_filename = "relative_importance_chart.png"
    if os.path.exists(chart_filename):
        doc.add_heading("Relative Importance Chart", level=2)
        doc.add_picture(chart_filename, width=Inches(6))
    else:
        logger.warning(
            f"Chart file {chart_filename} not found. Skipping chart insertion."
        )

    # Add Business Analysis Section
    doc.add_heading("Business Analysis", level=1)
    add_markdown_to_docx(business_analysis.analysis, doc)

    # Add Recommendations Section
    doc.add_heading("Recommendations", level=1)
    add_markdown_to_docx(business_analysis.recommendations, doc)

    # Add Conclusion
    doc.add_heading("Conclusion", level=1)
    add_markdown_to_docx(business_analysis.conclusions, doc)

    # Add Appendix Section: Statements grouped by sentiment
    doc.add_heading("Appendix: Statements by Sentiment", level=1)

    # Iterate over primary attributes and create the appendix table
    appendix_table = doc.add_table(rows=1, cols=4)
    appendix_table.style = "Light List Accent 1"

    # Set the header cells
    appendix_hdr_cells = appendix_table.rows[0].cells
    appendix_hdr_cells[0].text = "Primary Attribute"
    appendix_hdr_cells[1].text = "Positive Statements"
    appendix_hdr_cells[2].text = "Neutral Statements"
    appendix_hdr_cells[3].text = "Negative Statements"

    # Group statements by primary attribute and sentiment
    for primary_attr in customer_attributes:
        positive_statements = []
        neutral_statements = []
        negative_statements = []

        for secondary_attr in primary_attr.secondary_attributes:
            for tertiary_attr in secondary_attr.tertiary_attributes:
                for stmt in tertiary_attr.statements:
                    sentiment = stmt.sentiment_score
                    if sentiment > 0:
                        positive_statements.append(
                            f"{stmt.statement} ({sentiment:.2f})"
                        )
                    elif sentiment == 0:
                        neutral_statements.append(f"{stmt.statement} ({sentiment:.2f})")
                    else:
                        negative_statements.append(
                            f"{stmt.statement} ({sentiment:.2f})"
                        )

        # Add a new row for the primary attribute
        row_cells = appendix_table.add_row().cells
        row_cells[0].text = primary_attr.primary_attribute
        row_cells[1].text = (
            "\n".join(positive_statements) if positive_statements else "-"
        )
        row_cells[2].text = "\n".join(neutral_statements) if neutral_statements else "-"
        row_cells[3].text = (
            "\n".join(negative_statements) if negative_statements else "-"
        )

    # Save DOCX file
    report_filename = "final_customer_feedback_report.docx"
    try:
        doc.save(report_filename)
        logger.info(f"Final report saved to {report_filename}")
    except Exception as e:
        logger.error(f"An error occurred while saving the report: {e}")


# -----------------------------------------------------
# Saving Attributes JSON for Manual Review
# -----------------------------------------------------
def save_attributes_json(
    customer_attributes: List[PrimaryAttribute], file_path: str = "./attributes.json"
) -> None:
    """
    Save the derived customer attributes to a JSON file for manual review.

    Args:
        customer_attributes (List[PrimaryAttribute]): List of customer attributes.
        file_path (str): Path to save the JSON file.
    """
    logger.info(f"Saving customer attributes to {file_path} for manual review...")
    try:
        # Serialize the Pydantic model to JSON using model_dump_json()
        attributes_json = CustomerAttributes(
            attributes=customer_attributes
        ).model_dump_json(indent=4)
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(attributes_json)
        logger.info(f"Customer attributes successfully saved to {file_path}")
    except Exception as e:
        logger.error(f"Failed to save customer attributes to {file_path}: {e}")


# -----------------------------------------------------
# Main Function
# -----------------------------------------------------
def main():
    """
    Main function to execute the customer attribute analysis pipeline.
    """
    # Parse command-line arguments
    args = parse_arguments()
    domain = args.domain
    input_file = args.input_file
    log_level = args.log_level.upper()

    # Adjust logging level if needed
    if log_level in ["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"]:
        logger.setLevel(getattr(logging, log_level))
        for handler in logger.handlers:
            handler.setLevel(getattr(logging, log_level))
    else:
        logger.warning(f"Invalid log level '{log_level}'. Defaulting to INFO.")
        logger.setLevel(logging.INFO)
        for handler in logger.handlers:
            handler.setLevel(logging.INFO)

    logger.info("Starting the customer attribute analysis pipeline...")

    # Load customer statements
    customer_statements = load_customer_statements(input_file)

    if not customer_statements:
        logger.error("No customer statements loaded. Exiting.")
        return

    # Create LLM chains
    sentiment_sequence = create_sentiment_chain()
    attribute_sequence = create_attribute_chain()

    # Step 1: Run sentiment analysis chain
    logger.info("Performing sentiment analysis...")

    # Perform sentiment analysis with error handling
    try:
        raw_output = sentiment_sequence.invoke(
            {"statements": "\n".join(customer_statements), "domain": domain}
        )
        logger.debug(f"Raw LLM output: {raw_output}")

        # Check if results are present
        if not raw_output or not raw_output.results:
            logger.error("No sentiment results obtained. Exiting.")
            return

        parsed_sentiment_results: List[SentimentResult] = raw_output.results
        logger.debug(f"Parsed Sentiment Results: {parsed_sentiment_results}")

    except OutputParserException as e:
        logger.error(f"Failed to parse sentiment results: {e}")
        return
    except Exception as e:
        logger.error(f"An unexpected error occurred during sentiment analysis: {e}")
        return

    # Step 2: Run attribute derivation chain
    logger.info("Deriving customer attributes...")
    try:
        customer_attributes_container = attribute_sequence.invoke(
            {
                "sentiment_results": [
                    result.dict() for result in parsed_sentiment_results
                ],
                "domain": domain,
            }
        )

        if (
            not customer_attributes_container
            or not customer_attributes_container.attributes
        ):
            logger.error("No customer attributes derived. Exiting.")
            return

        parsed_customer_attributes: List[
            PrimaryAttribute
        ] = customer_attributes_container.attributes
        logger.debug(f"Parsed Customer Attributes: {parsed_customer_attributes}")

        # Save attributes.json for manual review
        save_attributes_json(parsed_customer_attributes)

    except OutputParserException as e:
        logger.error(f"Failed to parse customer attributes: {e}")
        return
    except Exception as e:
        logger.error(f"An unexpected error occurred during attribute derivation: {e}")
        return

    # Step 3: Calculate relative importance
    attribute_weights = calculate_relative_importance(parsed_customer_attributes)
    logger.debug(f"Attribute Weights before Conversion: {attribute_weights}")

    if not attribute_weights:
        logger.error("No attribute weights calculated. Exiting.")
        return

    # Step 4: Generate visualizations
    generate_visualizations(attribute_weights)

    # Step 5: Generate business analysis
    survey_details = load_survey_details("./survey_details.txt")
    # pdb.set_trace()

    business_analysis = generate_business_analysis(
        domain, parsed_customer_attributes, attribute_weights, survey_details
    )

    # Step 6: Generate the comprehensive report with structured data and business analysis
    # pdb.set_trace()
    generate_final_report(
        domain, parsed_customer_attributes, attribute_weights, business_analysis
    )


# -----------------------------------------------------
# Entry Point
# -----------------------------------------------------
if __name__ == "__main__":
    main()
